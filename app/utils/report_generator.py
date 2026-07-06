"""
Turnitin-style PDF report generator.

NEW APPROACH: Preserves the original uploaded PDF and overlays highlight
annotations directly on it. Summary/overview pages are prepended.

Produces three report types:
  1. Plagiarism Similarity Report (light brown highlights on original PDF)
  2. AI Writing Detection Report (sky blue highlights on original PDF)
  3. Combined Originality Report (both)
"""

import io
import os
import re
import base64
import logging
import fitz  # PyMuPDF
import httpx
from jinja2 import Environment, FileSystemLoader
from markupsafe import Markup
from xhtml2pdf import pisa
from app.models.document import ScanDocument

logger = logging.getLogger(__name__)

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")
ASSETS_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "assets")


def _get_logo_base64() -> str:
    """Load the brand logo and return as base64 data URI for HTML embedding."""
    logo_path = os.path.join(ASSETS_DIR, "image.png")
    try:
        with open(logo_path, "rb") as f:
            logo_bytes = f.read()
        b64 = base64.b64encode(logo_bytes).decode("utf-8")
        return f"data:image/png;base64,{b64}"
    except Exception as e:
        logger.warning(f"Could not load logo from {logo_path}: {e}")
        return ""


def _icon_to_data_uri(draw_fn, size: int = 64) -> str:
    """Render an icon as PNG using PyMuPDF and return as base64 data URI.
    xhtml2pdf cannot render SVG — we draw shapes via PyMuPDF then export PNG."""
    try:
        doc = fitz.open()
        page = doc.new_page(width=size, height=size)
        shape = page.new_shape()
        draw_fn(shape, size)
        shape.commit()
        pix = page.get_pixmap(alpha=True)
        png = pix.tobytes("png")
        doc.close()
        return f"data:image/png;base64,{base64.b64encode(png).decode()}"
    except Exception as e:
        logger.warning(f"Icon generation failed: {e}")
        return ""


def _get_report_icons() -> dict:
    """Generate all report icons as base64 PNG data URIs via PyMuPDF.

    Match Groups  — colored filled circles with white symbols
    Top Sources   — gray outlined icons (globe, newspaper, person)
    AI Detection  — colored filled circles with white symbols
    """
    S = 64  # Canvas size (points = pixels at 72 DPI)
    WHITE = (1, 1, 1)

    # ── Match Group 1: Not Cited or Quoted — RED circle, white doc ──
    def _not_cited(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.878, 0.322, 0.322), color=None)
        sh.draw_rect(fitz.Rect(s*0.30, s*0.20, s*0.70, s*0.72))
        sh.finish(fill=WHITE, color=None)
        for y in [0.34, 0.44, 0.54]:
            sh.draw_line(fitz.Point(s*0.38, s*y), fitz.Point(s*0.62, s*y))
            sh.finish(color=(0.878, 0.322, 0.322), width=2)

    # ── Match Group 2: Missing Quotations — ORANGE circle, white bubbles ──
    def _missing_quote(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.910, 0.569, 0.227), color=None)
        sh.draw_rect(fitz.Rect(s*0.18, s*0.20, s*0.55, s*0.50))
        sh.finish(fill=WHITE, color=None)
        sh.draw_rect(fitz.Rect(s*0.40, s*0.40, s*0.80, s*0.70))
        sh.finish(fill=WHITE, color=None)

    # ── Match Group 3: Missing Citation — OLIVE circle, white lines ──
    def _missing_citation(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.722, 0.659, 0.541), color=None)
        for i, y in enumerate([0.35, 0.50, 0.65]):
            w = s * 0.50 if i < 2 else s * 0.36
            sh.draw_line(fitz.Point(s*0.25, s*y), fitz.Point(s*0.25 + w, s*y))
            sh.finish(color=WHITE, width=3)

    # ── Match Group 4: Cited and Quoted — TEAL circle, white checkmark ──
    def _cited_quoted(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.165, 0.749, 0.671), color=None)
        pts = [fitz.Point(s*0.25, s*0.52),
               fitz.Point(s*0.42, s*0.68),
               fitz.Point(s*0.75, s*0.34)]
        sh.draw_polyline(pts)
        sh.finish(color=WHITE, width=4, closePath=False)

    # ── Top Sources: Internet — gray globe ──
    def _internet(sh, s):
        g = (0.33, 0.33, 0.33)
        sh.draw_circle(fitz.Point(s/2, s/2), s*0.42)
        sh.finish(color=g, width=1.8)
        sh.draw_line(fitz.Point(s*0.08, s/2), fitz.Point(s*0.92, s/2))
        sh.finish(color=g, width=1.2)
        sh.draw_line(fitz.Point(s/2, s*0.08), fitz.Point(s/2, s*0.92))
        sh.finish(color=g, width=1.2)
        sh.draw_line(fitz.Point(s*0.18, s*0.30), fitz.Point(s*0.82, s*0.30))
        sh.finish(color=g, width=0.8)
        sh.draw_line(fitz.Point(s*0.18, s*0.70), fitz.Point(s*0.82, s*0.70))
        sh.finish(color=g, width=0.8)

    # ── Top Sources: Publications — gray newspaper ──
    def _publications(sh, s):
        g = (0.33, 0.33, 0.33)
        sh.draw_rect(fitz.Rect(s*0.12, s*0.12, s*0.88, s*0.88))
        sh.finish(color=g, width=1.8)
        sh.draw_rect(fitz.Rect(s*0.22, s*0.22, s*0.50, s*0.48))
        sh.finish(fill=g, color=None)
        sh.draw_line(fitz.Point(s*0.56, s*0.28), fitz.Point(s*0.80, s*0.28))
        sh.finish(color=g, width=1.5)
        sh.draw_line(fitz.Point(s*0.56, s*0.42), fitz.Point(s*0.76, s*0.42))
        sh.finish(color=g, width=1.5)
        sh.draw_line(fitz.Point(s*0.22, s*0.60), fitz.Point(s*0.80, s*0.60))
        sh.finish(color=g, width=1.2)
        sh.draw_line(fitz.Point(s*0.22, s*0.72), fitz.Point(s*0.70, s*0.72))
        sh.finish(color=g, width=1.2)

    # ── Top Sources: Student Papers — gray person ──
    def _student(sh, s):
        g = (0.33, 0.33, 0.33)
        sh.draw_circle(fitz.Point(s/2, s*0.30), s*0.16)
        sh.finish(color=g, width=2)
        pts = [fitz.Point(s*0.10, s*0.92),
               fitz.Point(s*0.28, s*0.56),
               fitz.Point(s*0.72, s*0.56),
               fitz.Point(s*0.90, s*0.92)]
        sh.draw_polyline(pts)
        sh.finish(color=g, width=2, closePath=False)

    # ── AI Detection: Generated — GREEN circle, white face ──
    def _ai_gen(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.0, 0.722, 0.580), color=None)
        sh.draw_circle(fitz.Point(s/2, s*0.44), s*0.22)
        sh.finish(color=WHITE, width=2)
        sh.draw_circle(fitz.Point(s*0.38, s*0.40), s*0.04)
        sh.finish(fill=WHITE, color=None)
        sh.draw_circle(fitz.Point(s*0.62, s*0.40), s*0.04)
        sh.finish(fill=WHITE, color=None)
        sh.draw_line(fitz.Point(s*0.38, s*0.54), fitz.Point(s*0.62, s*0.54))
        sh.finish(color=WHITE, width=2)
        sh.draw_line(fitz.Point(s*0.50, s*0.22), fitz.Point(s*0.50, s*0.14))
        sh.finish(color=WHITE, width=2)
        sh.draw_rect(fitz.Rect(s*0.32, s*0.70, s*0.68, s*0.84))
        sh.finish(color=WHITE, width=1.5)

    # ── AI Detection: Paraphrased — BLUE circle, white face + wave ──
    def _ai_para(sh, s):
        sh.draw_circle(fitz.Point(s/2, s/2), s/2 - 2)
        sh.finish(fill=(0.035, 0.518, 0.890), color=None)
        sh.draw_circle(fitz.Point(s/2, s*0.40), s*0.22)
        sh.finish(color=WHITE, width=2)
        sh.draw_circle(fitz.Point(s*0.38, s*0.36), s*0.04)
        sh.finish(fill=WHITE, color=None)
        sh.draw_circle(fitz.Point(s*0.62, s*0.36), s*0.04)
        sh.finish(fill=WHITE, color=None)
        sh.draw_line(fitz.Point(s*0.38, s*0.50), fitz.Point(s*0.62, s*0.50))
        sh.finish(color=WHITE, width=2)
        sh.draw_line(fitz.Point(s*0.50, s*0.18), fitz.Point(s*0.50, s*0.10))
        sh.finish(color=WHITE, width=2)
        pts = [fitz.Point(s*0.20, s*0.78), fitz.Point(s*0.38, s*0.70),
               fitz.Point(s*0.55, s*0.82), fitz.Point(s*0.80, s*0.70)]
        sh.draw_polyline(pts)
        sh.finish(color=WHITE, width=2, closePath=False)

    return {
        "icon_not_cited": _icon_to_data_uri(_not_cited, S),
        "icon_missing_quote": _icon_to_data_uri(_missing_quote, S),
        "icon_missing_citation": _icon_to_data_uri(_missing_citation, S),
        "icon_cited_quoted": _icon_to_data_uri(_cited_quoted, S),
        "icon_internet": _icon_to_data_uri(_internet, S),
        "icon_publications": _icon_to_data_uri(_publications, S),
        "icon_student": _icon_to_data_uri(_student, S),
        "icon_ai_gen": _icon_to_data_uri(_ai_gen, S),
        "icon_ai_para": _icon_to_data_uri(_ai_para, S),
    }


# Highlight colors (R, G, B) — values 0-1
COLOR_PLAGIARISM = (0.96, 0.82, 0.66)  # Light brown #f5d0a9
COLOR_AI = (0.75, 0.86, 0.99)          # Sky blue #bfdbfe


def _get_template(name: str):
    env = Environment(loader=FileSystemLoader(TEMPLATE_DIR))
    return env.get_template(name)


def _html_to_pdf_bytes(html: str) -> bytes:
    """Convert HTML string to PDF bytes using xhtml2pdf."""
    buf = io.BytesIO()
    status = pisa.CreatePDF(html, dest=buf)
    if status.err:
        raise Exception("Failed to convert HTML to PDF.")
    pdf = buf.getvalue()
    buf.close()
    return pdf


def _download_original_file(url: str) -> bytes:
    """Download the original file from Cloudinary."""
    logger.info(f"Downloading original file from: {url[:80]}...")
    resp = httpx.get(url, follow_redirects=True, timeout=30)
    resp.raise_for_status()
    logger.info(f"Downloaded {len(resp.content)} bytes")
    return resp.content


def _is_pdf_file(doc: ScanDocument) -> bool:
    """Check if the original uploaded file is a PDF."""
    file_type = (doc.file_type or "").lower()
    file_name = (doc.original_file_name or "").lower()
    return file_type == "pdf" or file_name.endswith(".pdf")


def _is_docx_file(doc: ScanDocument) -> bool:
    """Check if the original uploaded file is a DOCX."""
    file_type = (doc.file_type or "").lower()
    file_name = (doc.original_file_name or "").lower()
    return file_type == "docx" or file_name.endswith(".docx")


def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert DOCX file bytes into PDF bytes.
    
    Uses python-docx to read the document structure and xhtml2pdf to
    render it as a clean PDF. Preserves:
      - Paragraphs and line breaks
      - Headings (bold, larger font)
      - Tables
      - Basic text formatting (bold, italic, underline)
    
    This is an ISOLATED function — it does NOT touch any other code path.
    """
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from html import escape

    logger.info("Converting DOCX to PDF...")
    doc = DocxDocument(io.BytesIO(docx_bytes))

    html_parts = [
        '<!DOCTYPE html>',
        '<html><head><meta charset="utf-8"/>',
        '<style>',
        '  @page { size: A4; margin: 2cm; }',
        '  body { font-family: "Times New Roman", Times, serif; font-size: 12pt; line-height: 1.5; color: #222; }',
        '  h1 { font-size: 18pt; font-weight: bold; margin: 12pt 0 6pt 0; }',
        '  h2 { font-size: 15pt; font-weight: bold; margin: 10pt 0 5pt 0; }',
        '  h3 { font-size: 13pt; font-weight: bold; margin: 8pt 0 4pt 0; }',
        '  p { margin: 3pt 0; text-align: justify; }',
        '  p.center { text-align: center; }',
        '  p.right { text-align: right; }',
        '  table { border-collapse: collapse; width: 100%; margin: 8pt 0; }',
        '  td, th { border: 1px solid #999; padding: 4pt 6pt; font-size: 11pt; }',
        '  th { background: #f0f0f0; font-weight: bold; }',
        '</style>',
        '</head><body>',
    ]

    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            html_parts.append('<p>&nbsp;</p>')
            continue

        # Determine heading level
        style_name = (para.style.name or "").lower()
        if 'heading 1' in style_name:
            tag = 'h1'
        elif 'heading 2' in style_name:
            tag = 'h2'
        elif 'heading 3' in style_name or 'heading' in style_name:
            tag = 'h3'
        else:
            tag = 'p'

        # Alignment
        align_class = ''
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            align_class = ' class="center"'
        elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            align_class = ' class="right"'

        # Build run-level formatting
        run_html = ''
        for run in para.runs:
            chunk = escape(run.text)
            if not chunk:
                continue
            if run.bold:
                chunk = f'<b>{chunk}</b>'
            if run.italic:
                chunk = f'<i>{chunk}</i>'
            if run.underline:
                chunk = f'<u>{chunk}</u>'
            run_html += chunk

        # Fallback if no runs extracted
        if not run_html:
            run_html = escape(text)

        html_parts.append(f'<{tag}{align_class}>{run_html}</{tag}>')

    # Process tables
    for table in doc.tables:
        html_parts.append('<table>')
        for i, row in enumerate(table.rows):
            html_parts.append('<tr>')
            cell_tag = 'th' if i == 0 else 'td'
            for cell in row.cells:
                html_parts.append(f'<{cell_tag}>{escape(cell.text)}</{cell_tag}>')
            html_parts.append('</tr>')
        html_parts.append('</table>')

    html_parts.append('</body></html>')
    html_str = '\n'.join(html_parts)

    pdf_bytes = _html_to_pdf_bytes(html_str)
    logger.info(f"DOCX → PDF conversion complete: {len(pdf_bytes)} bytes")
    return pdf_bytes


def _get_original_as_pdf(doc: ScanDocument) -> bytes | None:
    """
    Download the original file and return it as PDF bytes.
    - PDF files → returned as-is
    - DOCX files → converted to PDF first
    - Other files → returns None (unsupported)
    
    This is the SINGLE entry point for getting a usable PDF from the original.
    """
    if not doc.original_file_url:
        return None

    try:
        raw_bytes = _download_original_file(doc.original_file_url)

        if _is_pdf_file(doc):
            logger.info("Original is PDF, using directly")
            return raw_bytes
        elif _is_docx_file(doc):
            logger.info("Original is DOCX, converting to PDF")
            return _convert_docx_to_pdf(raw_bytes)
        else:
            logger.warning(f"Unsupported file type: {doc.file_type}")
            return None
    except Exception as e:
        logger.error(f"Failed to get original as PDF: {e}")
        import traceback
        traceback.print_exc()
        return None


def _highlight_text_in_pdf(pdf_bytes: bytes, texts_to_highlight: list[str], color: tuple) -> bytes:
    """
    Open a PDF, search for each text snippet, and add colored highlight
    annotations at the exact positions. Returns the modified PDF bytes.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for text in texts_to_highlight:
        if not text or len(text.strip()) < 10:
            continue

        # Search across all pages
        for page in doc:
            # Search for the text — returns list of Rect objects
            text_instances = page.search_for(text, quads=True)
            if text_instances:
                for quad in text_instances:
                    annot = page.add_highlight_annot(quad)
                    annot.set_colors(stroke=color)
                    annot.set_opacity(0.4)
                    annot.update()

    result = doc.tobytes()
    doc.close()
    return result


def _highlight_sentences_in_pdf(pdf_bytes: bytes, sentences: list[str], color: tuple) -> bytes:
    """
    Highlight individual sentences in the PDF. For long sentences, search
    for the first 60 chars to increase match probability.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")

    for sentence in sentences:
        if not sentence or len(sentence.strip()) < 10:
            continue

        # Use shorter search strings for better matching
        # (PDF text extraction might have different whitespace)
        search_text = sentence.strip()
        if len(search_text) > 80:
            search_text = search_text[:80]

        for page in doc:
            text_instances = page.search_for(search_text, quads=True)
            if text_instances:
                for quad in text_instances:
                    annot = page.add_highlight_annot(quad)
                    annot.set_colors(stroke=color)
                    annot.set_opacity(0.4)
                    annot.update()
                break  # Found on this page, move to next sentence

    result = doc.tobytes()
    doc.close()
    return result


def _merge_pdfs(summary_pdf_bytes: bytes, original_pdf_bytes: bytes) -> bytes:
    """Prepend the summary pages before the original document pages."""
    summary_doc = fitz.open(stream=summary_pdf_bytes, filetype="pdf")
    original_doc = fitz.open(stream=original_pdf_bytes, filetype="pdf")

    # Insert summary pages at the beginning of the original
    original_doc.insert_pdf(summary_doc, to_page=-1, start_at=0)

    result = original_doc.tobytes()
    summary_doc.close()
    original_doc.close()
    return result


# ────────────────────────────────────────────────────────────────────────────
#  Collect texts to highlight
# ────────────────────────────────────────────────────────────────────────────

def _get_plagiarism_texts(doc: ScanDocument) -> list[str]:
    """Collect plagiarized text snippets from scan results.
    Uses matched_text from LLM analysis + chunk text with dynamic threshold."""
    texts = []
    seen = set()
    plag_score = doc.plagiarism_result.plagiarism_score if doc.plagiarism_result else 0

    # Dynamic threshold: lower threshold for high-plagiarism docs to catch more
    threshold = max(5, min(20, 30 - plag_score * 0.3))

    if doc.plagiarism_result and doc.plagiarism_result.chunks:
        for chunk in doc.plagiarism_result.chunks:
            if chunk.plagiarism_score >= threshold and chunk.text:
                # Add matched_text from sources (more precise than whole chunk)
                for src_info in chunk.sources:
                    if isinstance(src_info, dict):
                        mt = src_info.get("matched_text", "")
                        if mt and len(mt.strip()) >= 20 and mt not in seen:
                            seen.add(mt)
                            texts.append(mt)

                # Also add the whole chunk text if score is high
                if chunk.plagiarism_score >= threshold * 1.5 and chunk.text not in seen:
                    seen.add(chunk.text)
                    texts.append(chunk.text)

    if doc.plagiarism_result and doc.plagiarism_result.matched_sources:
        for source in doc.plagiarism_result.matched_sources:
            if source.similarity_score >= threshold:
                if source.matched_text and source.matched_text not in seen:
                    seen.add(source.matched_text)
                    texts.append(source.matched_text)
                if source.original_text and source.original_text not in seen:
                    seen.add(source.original_text)
                    texts.append(source.original_text)

    return texts


def _get_ai_sentences(doc: ScanDocument) -> list[str]:
    """
    Determine which sentences to highlight as AI-generated using
    multi-signal per-sentence analysis — NOT just keyword matching.

    Signals per sentence:
    1. Length uniformity — AI sentences cluster around 15-25 words
    2. Vocabulary monotony — low type-token ratio within sentence
    3. AI phrase markers — known LLM filler phrases
    4. Transition smoothness — AI over-uses smooth connectors
    5. Neighbor similarity — AI sentences have similar length to neighbors
    """
    text = doc.extracted_text
    ai_score = doc.ai_result.ai_score if doc.ai_result else 0

    if ai_score <= 5 or not text:
        return []

    # Split into sentences
    sentence_pattern = re.compile(r'[^.!?]+[.!?]+')
    sentences = [m.group().strip() for m in sentence_pattern.finditer(text)]

    if not sentences:
        return []

    # Compute word counts per sentence for neighbor analysis
    word_counts = [len(s.split()) for s in sentences]
    avg_wc = sum(word_counts) / len(word_counts) if word_counts else 20

    ai_keywords = [
        "delve", "tapestry", "moreover", "furthermore", "testament", "notably",
        "in conclusion", "it is important to note", "consequently", "pivotal",
        "beacon", "comprehensive", "demystify", "multifaceted", "paramount",
        "additionally", "in this context", "in the realm of", "is crucial",
        "when it comes to", "needless to say", "it goes without saying",
        "it's worth noting", "having said that", "on the other hand",
        "in summary", "to summarize", "as a result",
    ]

    smooth_starters = (
        "this ", "these ", "those ", "such ", "the ", "in ", "as ",
        "however,", "therefore,", "consequently,", "furthermore,",
        "additionally,", "moreover,", "thus,",
    )

    scored = []
    for i, sent in enumerate(sentences):
        sent_lower = sent.lower().strip()
        wc = word_counts[i]
        words = re.findall(r'\b\w+\b', sent_lower)
        score = 0.0

        # Signal 1: Length uniformity (AI clusters around 15-25 words)
        if 14 <= wc <= 26:
            score += 3.0
        elif 10 <= wc <= 30:
            score += 1.0

        # Signal 2: Sentence-level vocabulary monotony
        if len(words) >= 5:
            unique_ratio = len(set(words)) / len(words)
            if unique_ratio < 0.55:
                score += 4.0  # Very repetitive vocabulary
            elif unique_ratio < 0.65:
                score += 2.0

        # Signal 3: AI phrase markers (weighted by specificity)
        for kw in ai_keywords:
            if kw in sent_lower:
                score += 8.0
                break  # One match is enough signal

        # Signal 4: Smooth transition starters
        if sent_lower.startswith(smooth_starters):
            score += 2.0

        # Signal 5: Neighbor length similarity (AI = uniform length)
        if i > 0 and i < len(sentences) - 1:
            prev_wc = word_counts[i - 1]
            next_wc = word_counts[i + 1]
            avg_neighbor = (prev_wc + next_wc) / 2
            if avg_neighbor > 0:
                length_diff = abs(wc - avg_neighbor) / avg_neighbor
                if length_diff < 0.15:  # Very similar to neighbors
                    score += 3.0
                elif length_diff < 0.25:
                    score += 1.5

        # Signal 6: Closeness to document average (AI = mean-hugging)
        if avg_wc > 0:
            dev_from_mean = abs(wc - avg_wc) / avg_wc
            if dev_from_mean < 0.10:
                score += 2.0

        scored.append((score, sent))

    # Dynamic threshold: highlight top N% based on actual AI score
    num_to_highlight = max(1, int(len(sentences) * (ai_score / 100.0)))
    scored.sort(key=lambda x: x[0], reverse=True)

    return [sent for _, sent in scored[:num_to_highlight]]


# ────────────────────────────────────────────────────────────────────────────
#  Summary page generators (HTML → PDF for prepending)
# ────────────────────────────────────────────────────────────────────────────

def _build_plagiarism_summary_pdf(doc: ScanDocument) -> bytes:
    """Generate the summary/overview pages for plagiarism report.
    All values computed dynamically from actual scan results — nothing hardcoded."""
    metadata = doc.metadata or {}
    scanned_at = (
        doc.scanned_at.strftime("%b %d, %Y, %I:%M %p UTC")
        if doc.scanned_at
        else doc.created_at.strftime("%b %d, %Y, %I:%M %p UTC")
    )

    plag_score = doc.plagiarism_result.plagiarism_score if doc.plagiarism_result else 0
    chunks = doc.plagiarism_result.chunks if doc.plagiarism_result else []

    # ── Dynamic match group computation from chunk data ──
    not_cited_chunks = []
    missing_quote_chunks = []
    missing_citation_chunks = []
    cited_quoted_chunks = []

    for chunk in chunks:
        if chunk.plagiarism_score < 10:
            continue
        # Extract match_type from chunk sources
        match_type = "not_cited"  # default
        for src_info in chunk.sources:
            if isinstance(src_info, dict) and src_info.get("match_type"):
                match_type = src_info["match_type"]
                break

        if match_type == "missing_quote":
            missing_quote_chunks.append(chunk)
        elif match_type == "missing_citation":
            missing_citation_chunks.append(chunk)
        elif match_type == "cited_quoted":
            cited_quoted_chunks.append(chunk)
        elif match_type != "original":
            not_cited_chunks.append(chunk)

    total_flagged = len(not_cited_chunks) + len(missing_quote_chunks) + len(missing_citation_chunks) + len(cited_quoted_chunks)

    # Compute percentages proportionally from actual data
    if total_flagged > 0 and plag_score > 0:
        not_cited_pct = round((len(not_cited_chunks) / total_flagged) * plag_score, 1)
        missing_quote_pct = round((len(missing_quote_chunks) / total_flagged) * plag_score, 1)
        missing_citation_pct = round((len(missing_citation_chunks) / total_flagged) * plag_score, 1)
        cited_quoted_pct = round((len(cited_quoted_chunks) / total_flagged) * plag_score, 1)
    else:
        not_cited_pct = round(plag_score, 1)
        missing_quote_pct = 0
        missing_citation_pct = 0
        cited_quoted_pct = 0

    # ── Dynamic source breakdown ──
    internet_score = 0.0
    student_score = 0.0
    if doc.plagiarism_result and doc.plagiarism_result.matched_sources:
        for src in doc.plagiarism_result.matched_sources:
            if src.url == "Submitted Work (Student Paper)":
                student_score += src.similarity_score
            else:
                internet_score += src.similarity_score
    total_sim = internet_score + student_score
    internet_pct = round((internet_score / total_sim) * plag_score, 1) if total_sim > 0 else 0
    student_pct = round((student_score / total_sim) * plag_score, 1) if total_sim > 0 else 0

    # ── Prepare matched sources list ──
    matched_sources = []
    if doc.plagiarism_result and doc.plagiarism_result.matched_sources:
        for src in doc.plagiarism_result.matched_sources:
            if src.similarity_score < 5:
                continue
            source_type = "student" if src.url == "Submitted Work (Student Paper)" else "internet"
            total_src_sim = sum(s.similarity_score for s in doc.plagiarism_result.matched_sources if s.similarity_score >= 5)
            sim_pct = round((src.similarity_score / total_src_sim) * plag_score, 1) if total_src_sim > 0 else 0
            matched_sources.append({
                "title": src.title or src.url,
                "url": src.url,
                "similarity_pct": max(sim_pct, 1) if sim_pct > 0 else "<1",
                "source_type": source_type,
                "raw_score": src.similarity_score,
            })
        matched_sources.sort(key=lambda x: x["raw_score"], reverse=True)
        matched_sources = matched_sources[:15]

    # ── Dynamic filtered sections ──
    filtered_sections = []
    text_lower = (doc.extracted_text or "").lower()
    if "bibliography" in text_lower or "references" in text_lower:
        filtered_sections.append("Bibliography")
    if "abstract" in text_lower:
        filtered_sections.append("Abstract")

    template = _get_template("plagiarism_report.html")
    html = template.render(
        document_id=str(doc.id),
        file_name=doc.original_file_name,
        file_type=doc.file_type.upper(),
        scanned_at=scanned_at,
        page_count=metadata.get("page_count", "-"),
        word_count=f"{metadata.get('token_count', 0):,}",
        char_count=f"{metadata.get('character_count', 0):,}",
        overall_plagiarism_score=round(plag_score, 1),
        filtered_sections=filtered_sections,
        integrity_flags=doc.integrity_flags or [],
        integrity_flag_count=len(doc.integrity_flags) if doc.integrity_flags else 0,
        matched_sources=matched_sources,
        highlighted_text=Markup(""),
        not_cited_count=len(not_cited_chunks),
        not_cited_pct=not_cited_pct,
        missing_quote_count=len(missing_quote_chunks),
        missing_quote_pct=missing_quote_pct,
        missing_citation_count=len(missing_citation_chunks),
        missing_citation_pct=missing_citation_pct,
        cited_quoted_count=len(cited_quoted_chunks),
        cited_quoted_pct=cited_quoted_pct,
        internet_pct=internet_pct,
        publication_pct=0,
        student_pct=student_pct,
        logo_base64=_get_logo_base64(),
        **_get_report_icons(),
    )

    return _html_to_pdf_bytes(html)


def _build_ai_summary_pdf(doc: ScanDocument) -> bytes:
    """Generate the summary/overview pages for AI report.
    AI generated vs paraphrased split computed from heuristics — not hardcoded."""
    metadata = doc.metadata or {}
    ai_score = doc.ai_result.ai_score if doc.ai_result else 0
    heuristics = doc.ai_result.heuristics if doc.ai_result else {}

    scanned_at = (
        doc.scanned_at.strftime("%b %d, %Y, %I:%M %p UTC")
        if doc.scanned_at
        else doc.created_at.strftime("%b %d, %Y, %I:%M %p UTC")
    )

    # Dynamic caution level based on score
    if ai_score >= 76:
        caution_level = "High confidence: AI-generated."
    elif ai_score >= 56:
        caution_level = "Caution: Likely AI-generated."
    elif ai_score >= 36:
        caution_level = "Caution: Review required."
    elif ai_score >= 16:
        caution_level = "Low confidence: Mostly human."
    else:
        caution_level = "No significant AI detected."

    # Dynamic AI generated vs paraphrased split based on heuristics
    # If TTR is low AND burstiness is low → mostly direct AI generation
    # If TTR is higher but burstiness is low → likely AI-paraphrased
    ttr = heuristics.get("type_token_ratio", 0.5)
    burstiness = heuristics.get("burstiness", 0.5)

    if ai_score > 0:
        # Compute paraphrase ratio: higher TTR with low burstiness = paraphrased
        paraphrase_signal = max(0, min(1, (ttr - 0.35) / 0.3))  # 0-1 scale
        uniformity_signal = max(0, min(1, (0.5 - burstiness) / 0.3))  # 0-1 scale

        paraphrase_ratio = paraphrase_signal * uniformity_signal * 0.4  # max 40% paraphrased
        paraphrase_ratio = max(0.05, min(0.40, paraphrase_ratio))  # clamp 5%-40%

        ai_paraphrased_pct = round(ai_score * paraphrase_ratio, 1)
        ai_generated_pct = round(ai_score - ai_paraphrased_pct, 1)
    else:
        ai_generated_pct = 0
        ai_paraphrased_pct = 0

    template = _get_template("ai_report.html")
    html = template.render(
        document_id=str(doc.id),
        file_name=doc.original_file_name,
        file_type=doc.file_type.upper(),
        scanned_at=scanned_at,
        page_count=metadata.get("page_count", "-"),
        word_count=f"{metadata.get('token_count', 0):,}",
        char_count=f"{metadata.get('character_count', 0):,}",
        overall_ai_score=round(ai_score, 1),
        caution_level=caution_level,
        ai_generated_pct=ai_generated_pct,
        ai_paraphrased_pct=ai_paraphrased_pct,
        burstiness=heuristics.get("burstiness"),
        type_token_ratio=heuristics.get("type_token_ratio"),
        avg_sentence_length=heuristics.get("avg_sentence_length"),
        ai_phrase_density=heuristics.get("ai_phrase_density"),
        highlighted_text=Markup(""),
        logo_base64=_get_logo_base64(),
        **_get_report_icons(),
    )

    return _html_to_pdf_bytes(html)


# ────────────────────────────────────────────────────────────────────────────
#  Public API: Build Report PDFs
# ────────────────────────────────────────────────────────────────────────────

def build_plagiarism_report_pdf(doc: ScanDocument) -> bytes:
    """
    Build Turnitin-style Plagiarism Report:
    1. Get original as PDF (auto-converts DOCX)
    2. Add light brown highlights on plagiarized text
    3. Generate summary pages
    4. Merge: summary + highlighted original
    """
    logger.info(f"Building plagiarism report for {doc.original_file_name}")

    # Step 1: Generate summary pages (always works)
    try:
        summary_pdf = _build_plagiarism_summary_pdf(doc)
        logger.info("Plagiarism summary PDF generated OK")
    except Exception as e:
        logger.error(f"Failed to build plagiarism summary: {e}", exc_info=True)
        raise

    # Step 2: Get original as PDF (handles PDF + DOCX)
    try:
        original_pdf = _get_original_as_pdf(doc)
    except Exception as e:
        logger.error(f"Failed to get original PDF: {e}", exc_info=True)
        original_pdf = None

    if not original_pdf:
        logger.warning("Could not get original as PDF, returning summary only")
        return summary_pdf

    # Step 3: Apply highlights
    try:
        plag_texts = _get_plagiarism_texts(doc)
        if plag_texts:
            original_pdf = _highlight_text_in_pdf(original_pdf, plag_texts, COLOR_PLAGIARISM)
    except Exception as e:
        logger.error(f"Failed to apply highlights: {e}", exc_info=True)

    # Step 4: Merge
    try:
        return _merge_pdfs(summary_pdf, original_pdf)
    except Exception as e:
        logger.error(f"Failed to merge PDFs: {e}", exc_info=True)
        return summary_pdf


def build_ai_report_pdf(doc: ScanDocument) -> bytes:
    """
    Build Turnitin-style AI Detection Report:
    1. Get original as PDF (auto-converts DOCX)
    2. Add sky blue highlights on AI-detected sentences
    3. Generate summary pages
    4. Merge: summary + highlighted original
    """
    logger.info(f"Building AI report for {doc.original_file_name}")

    # Generate summary pages (always works)
    summary_pdf = _build_ai_summary_pdf(doc)

    # Get original as PDF (handles PDF + DOCX)
    original_pdf = _get_original_as_pdf(doc)
    if not original_pdf:
        logger.warning("Could not get original as PDF, returning summary only")
        return summary_pdf

    # Apply highlights
    ai_sentences = _get_ai_sentences(doc)
    if ai_sentences:
        original_pdf = _highlight_sentences_in_pdf(original_pdf, ai_sentences, COLOR_AI)

    return _merge_pdfs(summary_pdf, original_pdf)


def build_report_pdf(doc: ScanDocument) -> bytes:
    """
    Build combined report (both plagiarism + AI highlights on original PDF).
    Backward compatible endpoint.
    """
    logger.info(f"Building combined report for {doc.original_file_name}")

    # Generate both summary pages (always works)
    plag_summary = _build_plagiarism_summary_pdf(doc)
    ai_summary = _build_ai_summary_pdf(doc)

    # Merge AI + Plagiarism summaries
    combined_summary = fitz.open(stream=ai_summary, filetype="pdf")
    plag_doc = fitz.open(stream=plag_summary, filetype="pdf")
    combined_summary.insert_pdf(plag_doc)
    summary_bytes = combined_summary.tobytes()
    combined_summary.close()
    plag_doc.close()

    # Get original as PDF (handles PDF + DOCX)
    original_pdf = _get_original_as_pdf(doc)
    if not original_pdf:
        logger.warning("Could not get original as PDF, returning summary only")
        return summary_bytes

    # Apply plagiarism highlights first (light brown)
    plag_texts = _get_plagiarism_texts(doc)
    if plag_texts:
        original_pdf = _highlight_text_in_pdf(original_pdf, plag_texts, COLOR_PLAGIARISM)

    # Then apply AI highlights (sky blue)
    ai_sentences = _get_ai_sentences(doc)
    if ai_sentences:
        original_pdf = _highlight_sentences_in_pdf(original_pdf, ai_sentences, COLOR_AI)

    return _merge_pdfs(summary_bytes, original_pdf)
